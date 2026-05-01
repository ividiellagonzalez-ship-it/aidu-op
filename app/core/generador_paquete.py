"""
AIDU Op · Generador de Paquete de Postulación
================================================
Genera Word (propuesta técnica + económica) y Excel (oferta económica)
con identidad AIDU Op aplicada.

Usa python-docx y openpyxl - ambos tienen wheels para Python 3.14.
"""
import logging
from pathlib import Path
from datetime import date, datetime
from typing import Dict, Optional

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from app.db.migrator import get_connection
from app.core.inteligencia_precios import calcular_escenarios_precio
from config.settings import AIDU_HOME

logger = logging.getLogger(__name__)

# Colores de identidad AIDU
AIDU_AZUL = RGBColor(0x1E, 0x40, 0xAF)
AIDU_AZUL_OSCURO = RGBColor(0x1E, 0x3A, 0x8A)
AIDU_GRIS = RGBColor(0x47, 0x55, 0x69)


def fmt_clp(n):
    if n is None:
        return "$0"
    return f"${int(n):,}".replace(",", ".")


def _set_cell_bg(cell, color_hex: str):
    """Pinta el fondo de una celda Word"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def generar_propuesta_tecnica(proyecto_id: int, output_dir: Path) -> Path:
    """Genera propuesta_tecnica.docx con identidad AIDU"""
    conn = get_connection()
    p = conn.execute("SELECT * FROM aidu_proyectos WHERE id = ?", (proyecto_id,)).fetchone()
    conn.close()
    if not p:
        raise ValueError(f"Proyecto {proyecto_id} no encontrado")

    doc = Document()

    # Estilo del documento
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # === PORTADA ===
    titulo_aidu = doc.add_paragraph()
    titulo_aidu.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_aidu = titulo_aidu.add_run('● AIDU')
    run_aidu.font.size = Pt(36)
    run_aidu.font.bold = True
    run_aidu.font.color.rgb = AIDU_AZUL
    run_op = titulo_aidu.add_run(' Op')
    run_op.font.size = Pt(20)
    run_op.font.color.rgb = AIDU_GRIS

    sub = doc.add_paragraph('Sistema de Gestión Comercial · Ingeniería · IA · Procesos')
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = AIDU_GRIS

    doc.add_paragraph()
    doc.add_paragraph()

    # Título principal
    titulo = doc.add_heading('PROPUESTA TÉCNICA', level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.runs[0].font.color.rgb = AIDU_AZUL

    # Datos de licitación en tabla
    tabla = doc.add_table(rows=4, cols=2)
    tabla.style = 'Light Grid Accent 1'
    datos = [
        ('Licitación:', p['nombre']),
        ('Código:', p['codigo_externo']),
        ('Organismo:', p['organismo'] or '-'),
        ('Fecha:', date.today().strftime('%d de %B de %Y')),
    ]
    for i, (k, v) in enumerate(datos):
        tabla.cell(i, 0).text = k
        tabla.cell(i, 1).text = str(v)
        tabla.cell(i, 0).paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # === SECCIÓN 1: ANTECEDENTES DEL OFERENTE ===
    h1 = doc.add_heading('1. Antecedentes del Oferente', level=1)
    h1.runs[0].font.color.rgb = AIDU_AZUL

    doc.add_paragraph(
        'AIDU Op SpA es una consultora chilena especializada en servicios profesionales '
        'de ingeniería, gestión de proyectos y transformación digital con inteligencia artificial. '
        'Liderada por Ignacio Vidiella González, Ingeniero Civil con especialización estructural '
        'y diplomados en Administración de Proyectos y Agilidad Organizacional.'
    )

    doc.add_paragraph(
        'Nuestra propuesta de valor combina rigor de ingeniería tradicional con herramientas '
        'modernas de inteligencia artificial y metodologías ágiles, lo que nos permite entregar '
        'soluciones de calidad superior en plazos competitivos.'
    )

    # === SECCIÓN 2: COMPRENSIÓN DEL SERVICIO ===
    h2 = doc.add_heading('2. Comprensión del Servicio', level=1)
    h2.runs[0].font.color.rgb = AIDU_AZUL

    doc.add_paragraph(
        f'El servicio licitado por {p["organismo"]} requiere lo siguiente:'
    )
    p_alcance = doc.add_paragraph()
    p_alcance.add_run('Alcance: ').bold = True
    p_alcance.add_run(p['descripcion'] or 'Ver bases técnicas')

    # === SECCIÓN 3: METODOLOGÍA TÉCNICA ===
    h3 = doc.add_heading('3. Metodología Técnica', level=1)
    h3.runs[0].font.color.rgb = AIDU_AZUL

    doc.add_paragraph('El servicio se ejecutará en cuatro fases:')

    fases = [
        ('Fase 1 — Levantamiento inicial',
         'Reunión de inicio con la contraparte técnica, revisión de antecedentes, '
         'definición de hitos y entregables, alineamiento de expectativas.'),
        ('Fase 2 — Análisis técnico',
         'Análisis técnico detallado conforme a las bases. Aplicación de normativa '
         'vigente y mejores prácticas de la disciplina.'),
        ('Fase 3 — Desarrollo de productos',
         'Elaboración de los entregables especificados con los estándares de calidad AIDU.'),
        ('Fase 4 — Revisión y entrega',
         'Revisión cruzada interna, entrega formal a la contraparte, sesión de '
         'transferencia de conocimiento y acompañamiento.'),
    ]
    for titulo_fase, contenido in fases:
        p_fase = doc.add_paragraph()
        p_fase.add_run(titulo_fase).bold = True
        p_fase.add_run(': ' + contenido)

    # === SECCIÓN 4: EQUIPO PROFESIONAL ===
    h4 = doc.add_heading('4. Equipo Profesional', level=1)
    h4.runs[0].font.color.rgb = AIDU_AZUL

    tabla_eq = doc.add_table(rows=3, cols=3)
    tabla_eq.style = 'Light Grid Accent 1'

    # Encabezado
    hdr = tabla_eq.rows[0].cells
    hdr[0].text = 'Rol'
    hdr[1].text = 'Profesional'
    hdr[2].text = 'HH dedicadas'
    for c in hdr:
        c.paragraphs[0].runs[0].bold = True
        _set_cell_bg(c, '1E40AF')
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    tabla_eq.cell(1, 0).text = 'Director Ejecutivo'
    tabla_eq.cell(1, 1).text = 'Ignacio Vidiella González · Ing. Civil'
    tabla_eq.cell(1, 2).text = f"{p['hh_ignacio_estimado'] or 0} h"

    tabla_eq.cell(2, 0).text = 'Socia Operacional'
    tabla_eq.cell(2, 1).text = 'Jorella · Ing. Comercial'
    tabla_eq.cell(2, 2).text = f"{p['hh_jorella_estimado'] or 0} h"

    # === SECCIÓN 5: EXPERIENCIA ===
    h5 = doc.add_heading('5. Experiencia AIDU', level=1)
    h5.runs[0].font.color.rgb = AIDU_AZUL

    doc.add_paragraph(
        'AIDU Op está construyendo track record en Mercado Público chileno. '
        'El director ejecutivo cuenta con experiencia previa en gestión operacional '
        'industrial chilena, incluyendo posiciones ejecutivas en empresas de ingeniería '
        'para minería de gran escala.'
    )

    # === FOOTER ===
    section = doc.sections[0]
    footer = section.footer
    p_footer = footer.paragraphs[0]
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = p_footer.add_run('AIDU Op SpA · Ignacio Vidiella González · RUT 16.402.949-2 · Machalí, O\'Higgins')
    run_f.font.size = Pt(9)
    run_f.font.color.rgb = AIDU_GRIS

    # Guardar
    output_path = output_dir / f"AIDU_Op_Propuesta_Tecnica_{p['codigo_externo']}.docx"
    doc.save(output_path)
    return output_path


def generar_oferta_economica_excel(proyecto_id: int, output_dir: Path) -> Path:
    """Genera oferta_economica.xlsx con tabla de partidas + fórmulas"""
    conn = get_connection()
    p = conn.execute("SELECT * FROM aidu_proyectos WHERE id = ?", (proyecto_id,)).fetchone()
    conn.close()
    if not p:
        raise ValueError(f"Proyecto {proyecto_id} no encontrado")

    # Calcular escenarios
    esc = calcular_escenarios_precio(proyecto_id)
    costo = esc["costo"]

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Oferta Económica"

    # Estilos AIDU
    header_fill = PatternFill(start_color="1E40AF", end_color="1E40AF", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    title_font = Font(bold=True, color="1E40AF", size=14)
    border_thin = Border(
        left=Side(style='thin', color='CBD5E1'),
        right=Side(style='thin', color='CBD5E1'),
        top=Side(style='thin', color='CBD5E1'),
        bottom=Side(style='thin', color='CBD5E1')
    )

    # === ENCABEZADO ===
    ws['A1'] = '● AIDU Op'
    ws['A1'].font = Font(bold=True, color="1E40AF", size=20)
    ws['A2'] = 'Sistema de Gestión Comercial'
    ws['A2'].font = Font(color="475569", size=10)

    ws['A4'] = 'OFERTA ECONÓMICA'
    ws['A4'].font = title_font

    # Datos licitación
    ws['A6'] = 'Licitación:'
    ws['B6'] = p['nombre']
    ws['A7'] = 'Código:'
    ws['B7'] = p['codigo_externo']
    ws['A8'] = 'Organismo:'
    ws['B8'] = p['organismo'] or '-'
    ws['A9'] = 'Fecha:'
    ws['B9'] = date.today().strftime('%d-%m-%Y')

    for row in range(6, 10):
        ws[f'A{row}'].font = Font(bold=True)

    # === TABLA DE PARTIDAS ===
    ws['A11'] = 'DESGLOSE DE COSTOS'
    ws['A11'].font = title_font

    headers = ['Partida', 'HH', 'Tarifa/h', 'Subtotal']
    for i, h in enumerate(headers, start=1):
        cell = ws.cell(row=12, column=i, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
        cell.border = border_thin

    # Filas de datos con fórmulas vivas
    hh_total = (p['hh_ignacio_estimado'] or 0) + (p['hh_jorella_estimado'] or 0)
    tarifa_hora = costo['tarifa_hora_clp']

    ws.cell(row=13, column=1, value='Honorarios profesionales')
    ws.cell(row=13, column=2, value=hh_total)
    ws.cell(row=13, column=3, value=tarifa_hora)
    ws.cell(row=13, column=4, value=f"=B13*C13")  # FÓRMULA

    ws.cell(row=14, column=1, value='Viajes y traslados')
    ws.cell(row=14, column=2, value='-')
    ws.cell(row=14, column=3, value='-')
    ws.cell(row=14, column=4, value=costo['viajes'])

    ws.cell(row=15, column=1, value='Subtotal')
    ws.cell(row=15, column=4, value=f"=D13+D14")  # FÓRMULA
    ws.cell(row=15, column=1).font = Font(bold=True)

    ws.cell(row=16, column=1, value='Overhead (18%)')
    ws.cell(row=16, column=4, value=f"=D15*0.18")  # FÓRMULA

    ws.cell(row=17, column=1, value='COSTO TOTAL')
    ws.cell(row=17, column=4, value=f"=D15+D16")  # FÓRMULA
    ws.cell(row=17, column=1).font = Font(bold=True, color="1E40AF")
    ws.cell(row=17, column=4).font = Font(bold=True, color="1E40AF")

    # Aplicar bordes y formato CLP
    for row in range(13, 18):
        for col in range(1, 5):
            cell = ws.cell(row=row, column=col)
            cell.border = border_thin
            if col == 4 and row != 13:  # columna Subtotal, excepto fila con HH
                cell.number_format = '"$"#,##0'

    # === ESCENARIOS DE PRECIO ===
    ws['A19'] = 'ESCENARIOS DE PRECIO (basados en histórico Mercado Público)'
    ws['A19'].font = title_font

    headers_esc = ['Escenario', 'Precio', 'Margen %', 'Probabilidad']
    for i, h in enumerate(headers_esc, start=1):
        cell = ws.cell(row=20, column=i, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
        cell.border = border_thin

    escenarios_data = [
        ('🥇 Agresivo', esc['agresivo']),
        ('⚡ Competitivo', esc['competitivo']),
        ('💎 Premium', esc['premium']),
    ]
    for i, (nombre, e) in enumerate(escenarios_data, start=21):
        ws.cell(row=i, column=1, value=nombre)
        ws.cell(row=i, column=2, value=e['precio'])
        ws.cell(row=i, column=2).number_format = '"$"#,##0'
        ws.cell(row=i, column=3, value=e['margen_pct'] / 100)
        ws.cell(row=i, column=3).number_format = '0.0%'
        ws.cell(row=i, column=4, value=e['probabilidad'] / 100)
        ws.cell(row=i, column=4).number_format = '0%'

        for col in range(1, 5):
            ws.cell(row=i, column=col).border = border_thin

    # === PRECIO FINAL OFERTADO ===
    ws['A26'] = 'PRECIO OFERTADO'
    ws['A26'].font = title_font

    ws['A27'] = 'Escenario elegido:'
    ws['B27'] = p['escenario_elegido'] or 'No definido'
    ws['A27'].font = Font(bold=True)

    ws['A28'] = 'Precio final:'
    ws['B28'] = p['precio_ofertado'] or esc['competitivo']['precio']
    ws['B28'].number_format = '"$"#,##0'
    ws['A28'].font = Font(bold=True, size=14, color="15803D")
    ws['B28'].font = Font(bold=True, size=14, color="15803D")

    # Anchos de columnas
    ws.column_dimensions['A'].width = 35
    ws.column_dimensions['B'].width = 25
    ws.column_dimensions['C'].width = 18
    ws.column_dimensions['D'].width = 18

    # Guardar
    output_path = output_dir / f"AIDU_Op_Oferta_Economica_{p['codigo_externo']}.xlsx"
    wb.save(output_path)
    return output_path


def generar_paquete_completo(proyecto_id: int) -> Dict:
    """
    Genera paquete completo: Word propuesta + Excel económica.
    Devuelve dict con paths de los archivos generados.
    """
    conn = get_connection()
    p = conn.execute("SELECT * FROM aidu_proyectos WHERE id = ?", (proyecto_id,)).fetchone()
    conn.close()
    if not p:
        raise ValueError(f"Proyecto {proyecto_id} no encontrado")

    # Carpeta del paquete
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    pkg_dir = AIDU_HOME / "paquetes" / f"AIDU_Op_{p['codigo_externo']}_{timestamp}"
    pkg_dir.mkdir(parents=True, exist_ok=True)

    archivos = {}

    try:
        archivos['propuesta_tecnica'] = generar_propuesta_tecnica(proyecto_id, pkg_dir)
        logger.info(f"✅ Propuesta técnica: {archivos['propuesta_tecnica']}")
    except Exception as e:
        logger.error(f"Error generando propuesta: {e}")

    try:
        archivos['oferta_economica'] = generar_oferta_economica_excel(proyecto_id, pkg_dir)
        logger.info(f"✅ Oferta económica: {archivos['oferta_economica']}")
    except Exception as e:
        logger.error(f"Error generando económica: {e}")

    # Marcar paquete generado en BD
    conn = get_connection()
    conn.execute(
        "UPDATE aidu_proyectos SET paquete_generado=1, paquete_path=?, fecha_modificacion=datetime('now','localtime') WHERE id=?",
        (str(pkg_dir), proyecto_id)
    )
    conn.commit()
    conn.close()

    return {
        "carpeta": pkg_dir,
        "archivos": archivos,
        "n_archivos": len(archivos),
    }
